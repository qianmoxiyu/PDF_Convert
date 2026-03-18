# -*- coding: utf-8 -*-
import tkinter as tk
from tkinter import filedialog, messagebox
import ttkbootstrap as ttk  # 替换标准ttk
from ttkbootstrap.constants import *
import fitz
import glob
import os
import re
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage, PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from docx import Document

class PDFConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF转换工具")
        self.root.geometry("500x400")
        
        # 使用ttkbootstrap主题（可选：'yeti, 'minty'等）
        self.style = ttk.Style(theme="yeti")  # 选用现代主题
        
        self.create_widgets()
    
    def create_widgets(self):
        # 主框架
        frame = ttk.Frame(self.root, padding=20)
        frame.pack(fill="both", expand=True)
        
        # 标题
        ttk.Label(frame, text="PDF转换工具", font=("Helvetica", 16, "bold")).pack(pady=10)
        
        # 操作按钮（使用更漂亮的按钮样式）
        ttk.Button(
            frame, 
            text="PDF转PNG", 
            command=self.convert_pdf_to_png,
            bootstyle=PRIMARY,# 蓝色主题按钮
            width=20
        ).pack(pady=10)
        
        ttk.Button(
            frame, 
            text="图片合并PDF", 
            command=self.merge_images_to_pdf,
            bootstyle=SUCCESS,  # 绿色主题按钮
            width=20
        ).pack(pady=10)
        
        ttk.Button(
            frame, 
            text="PDF转Word", 
            command=self.convert_pdf_to_word,
            bootstyle=INFO,  # 浅蓝色主题按钮
            width=20
        ).pack(pady=10)
        
        ttk.Button(
            frame, 
            text="退出", 
            command=self.root.quit,
            bootstyle=DANGER,  # 红色主题按钮
            width=20
        ).pack(pady=10)
        
        # 状态栏
        self.status_var = tk.StringVar()
        self.status_var.set("就绪")
        ttk.Label(frame, textvariable=self.status_var, bootstyle=INVERSE).pack(side="bottom", pady=10)
    
    def show_file_dialog(self, title, filetypes):
        file_path = filedialog.askopenfilename(title=title, filetypes=filetypes)
        return file_path if file_path else None
    
    def show_directory_dialog(self, title):
        dir_path = filedialog.askdirectory(title=title)
        return dir_path if dir_path else None
    
    def show_quality_dialog(self):
        quality = tk.StringVar(value="200%")
        
        dialog = tk.Toplevel(self.root)
        dialog.title("选择图片质量")
        
        ttk.Label(dialog, text="请选择图片质量:").pack(pady=10)
        
        options = ["200%缩放 (一般质量)", "300 DPI (高质量)", "600 DPI (超高质量)"]
        for option in options:
            ttk.Radiobutton(dialog, text=option, variable=quality, value=option).pack(anchor="w")
        
        def on_confirm():
            dialog.destroy()
        
        ttk.Button(dialog, text="确定", command=on_confirm).pack(pady=10)
        
        dialog.wait_window()
        return quality.get()
    
    def convert_pdf_to_png(self):
        self.status_var.set("正在执行PDF转PNG...")
        self.root.update()
        
        pdf_path = self.show_file_dialog("选择PDF文件", [("PDF文件", "*.pdf")])
        if not pdf_path:
            self.status_var.set("已取消")
            return
        
        output_dir = self.show_directory_dialog("选择PNG输出目录") or os.path.dirname(pdf_path)
        if not output_dir:
            self.status_var.set("已取消")
            return
        
        quality = self.show_quality_dialog()
        
        try:
            doc = fitz.open(pdf_path)
            total_pages = doc.page_count
            
            # 根据选择的质量设置参数
            if "300 DPI" in quality:
                mat = fitz.Matrix(300/72, 300/72)
            elif "600 DPI" in quality:
                mat = fitz.Matrix(600/72, 600/72)
            else:
                mat = fitz.Matrix(2.0, 2.0)
            
            for pg in range(total_pages):
                try:
                    page = doc[pg]
                    pm = page.get_pixmap(matrix=mat, alpha=False)
                    output_path = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(pdf_path))[0]}_{pg+1}.png")
                    pm.save(output_path)
                    
                    # 更新状态
                    self.status_var.set(f"正在转换第 {pg+1}/{total_pages} 页...")
                    self.root.update()
                
                except Exception as e:
                    messagebox.showerror("错误", f"页面 {pg+1} 转换失败: {str(e)}")
            
            doc.close()
            messagebox.showinfo("完成", f"PDF转PNG完成！共转换 {total_pages} 页。\n保存到: {output_dir}")
            self.status_var.set("PDF转PNG完成")
        
        except Exception as e:
            messagebox.showerror("错误", f"PDF处理失败: {str(e)}")
            self.status_var.set("转换失败")
    
    def merge_images_to_pdf(self):
        self.status_var.set("正在合并图片为PDF...")
        self.root.update()
        
        input_dir = self.show_directory_dialog("选择图片目录")
        if not input_dir:
            self.status_var.set("已取消")
            return
        
        output_dir = self.show_directory_dialog("选择PDF输出目录") or input_dir
        if not output_dir:
            self.status_var.set("已取消")
            return
        
        img_files = glob.glob(os.path.join(input_dir, "*.png"))
        if not img_files:
            messagebox.showerror("错误", f"目录中没有PNG图片: {input_dir}")
            self.status_var.set("没有找到图片")
            return
        
        # 自然排序
        def natural_sort(l):
            convert = lambda text: int(text) if text.isdigit() else text.lower()
            alphanum_key = lambda key: [convert(c) for c in re.split('([0-9]+)', key)]
            return sorted(l, key=alphanum_key)
        
        img_files = natural_sort(img_files)
        output_pdf = os.path.join(output_dir, f"{os.path.basename(input_dir)}_合并结果.pdf")
        
        try:
            doc = fitz.open()
            total_images = len(img_files)
            
            for i, img in enumerate(img_files):
                try:
                    imgdoc = fitz.open(img)
                    pdfbytes = imgdoc.convert_to_pdf()
                    imgpdf = fitz.open("pdf", pdfbytes)
                    doc.insert_pdf(imgpdf)
                    imgpdf.close()
                    
                    # 更新状态
                    self.status_var.set(f"正在处理第 {i+1}/{total_images} 张图片...")
                    self.root.update()
                
                except Exception as e:
                    messagebox.showerror("错误", f"处理图片失败 {os.path.basename(img)}: {str(e)}")
            
            if doc.page_count > 0:
                doc.save(output_pdf)
                messagebox.showinfo("完成", f"图片合并完成！共合并 {total_images} 张图片。\n保存到: {output_pdf}")
                self.status_var.set("图片合并完成")
            else:
                messagebox.showerror("错误", "没有有效的图片被添加")
                self.status_var.set("合并失败")
            
            doc.close()
        
        except Exception as e:
            messagebox.showerror("错误", f"合并失败: {str(e)}")
            self.status_var.set("合并失败")
    
    def convert_pdf_to_word(self):
        self.status_var.set("正在转换PDF到Word...")
        self.root.update()
        
        pdf_path = self.show_file_dialog("选择PDF文件", [("PDF文件", "*.pdf")])
        if not pdf_path:
            self.status_var.set("已取消")
            return
        
        output_dir = self.show_directory_dialog("选择Word输出目录") or os.path.dirname(pdf_path)
        if not output_dir:
            self.status_var.set("已取消")
            return
        
        try:
            document = Document()
            with open(pdf_path, 'rb') as fn:
                parser = PDFParser(fn)
                doc = PDFDocument(parser)
                parser.set_document(doc)
                doc.initialize("")
                
                if not doc.is_extractable:
                    raise PDFTextExtractionNotAllowed("PDF不允许文本提取")
                
                resource = PDFResourceManager()
                laparams = LAParams()
                device = PDFPageAggregator(resource, laparams=laparams)
                interpreter = PDFPageInterpreter(resource, device)
                
                pages = list(PDFPage.create_pages(doc))
                total_pages = len(pages)
                
                for i, page in enumerate(pages):
                    interpreter.process_page(page)
                    layout = device.get_result()
                    for out in layout:
                        if hasattr(out, "get_text"):
                            content = out.get_text().replace(u'\xa0', u' ')
                            document.add_paragraph(content, style='ListBullet')
                    
                    # 更新状态
                    self.status_var.set(f"正在转换第 {i+1}/{total_pages} 页...")
                    self.root.update()
            
            output_docx = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(pdf_path))[0]}_转换结果.docx")
            document.save(output_docx)
            messagebox.showinfo("完成", f"PDF转Word完成！共转换 {total_pages} 页。\n保存到: {output_docx}")
            self.status_var.set("PDF转Word完成")
        
        except Exception as e:
            messagebox.showerror("错误", f"转换失败: {str(e)}")
            self.status_var.set("转换失败")

if __name__ == '__main__':
    root = tk.Tk()
    app = PDFConverterApp(root)
    root.mainloop()